"""Convert LEGAL_CHECKLIST.md into a presentable .docx so the operator
can hand it to a lawyer or print it for filing.

This is intentionally a small script: it parses the markdown
structurally (headings, bullets, tables, emphasis) rather than via a
heavy converter, because the checklist's formatting is consistent and
we want full control over how the .docx looks.

Usage:
    python scripts/md_to_legal_docx.py LEGAL_CHECKLIST.md LEGAL_CHECKLIST.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# -------------------------------------------------------------------
# Inline markdown → docx runs
# -------------------------------------------------------------------

# We support **bold**, *italic*, `code`, and [text](url). Markdown's
# more exotic syntax (footnotes, tables-inside-cells, HTML) isn't
# used in the checklist so we don't bother handling it.
_INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*)"      # bold
    r"|(\*[^*]+\*)"          # italic
    r"|(`[^`]+`)"            # code
    r"|(\[[^\]]+\]\([^)]+\))"  # link
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Walk the text, splitting at the inline-markdown spans and
    adding a docx run per chunk with the right formatting."""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group()
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("[") and "](" in token:
            label, url = token[1:-1].split("](", 1)
            _add_hyperlink(paragraph, label, url)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a clickable hyperlink. python-docx exposes everything
    you need to build one, but doesn't have a one-liner — hence
    this XML-prodding."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # Word default link blue
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# -------------------------------------------------------------------
# Markdown block parser
# -------------------------------------------------------------------


def _parse_table_row(line: str) -> list[str]:
    # Strip leading + trailing pipes, split on |, strip cells.
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [c.strip() for c in inner.split("|")]


def _is_table_separator(line: str) -> bool:
    # Header separator looks like |---|---|---|. Allow alignment
    # markers like :--- or ---: or :---:.
    inner = line.strip().strip("|").strip()
    return all(re.fullmatch(r":?-+:?", c.strip()) for c in inner.split("|") if c.strip())


def iter_blocks(lines: list[str]) -> Iterator[dict]:
    """Group the source lines into logical blocks (heading, paragraph,
    bullet list, table, horizontal-rule). Tables are kept as a single
    block so the writer can build them as a docx table."""
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        if not line.strip():
            i += 1
            continue

        # Headings (#, ##, ###, ####)
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            yield {"type": "heading", "level": len(m.group(1)), "text": m.group(2)}
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", line.strip()):
            yield {"type": "hr"}
            i += 1
            continue

        # Table (two consecutive | lines, second being separator)
        if line.lstrip().startswith("|") and i + 1 < len(lines):
            sep = lines[i + 1].rstrip("\n")
            if _is_table_separator(sep):
                header = _parse_table_row(line)
                rows: list[list[str]] = []
                j = i + 2
                while j < len(lines) and lines[j].lstrip().startswith("|"):
                    rows.append(_parse_table_row(lines[j].rstrip("\n")))
                    j += 1
                yield {"type": "table", "header": header, "rows": rows}
                i = j
                continue

        # Bullet list
        if re.match(r"^[\s]*[-*]\s+", line):
            items: list[str] = []
            j = i
            while j < len(lines):
                cur = lines[j].rstrip("\n")
                m_bul = re.match(r"^[\s]*[-*]\s+(.*)$", cur)
                if m_bul:
                    items.append(m_bul.group(1))
                    j += 1
                elif cur.strip() and not re.match(r"^(#{1,4}\s|\||-{3,}$)", cur):
                    # Continuation of the previous bullet.
                    if items:
                        items[-1] = items[-1] + " " + cur.strip()
                    j += 1
                else:
                    break
            yield {"type": "bullets", "items": items}
            i = j
            continue

        # Paragraph — accumulate until blank line or new block.
        para_lines: list[str] = []
        while i < len(lines):
            cur = lines[i].rstrip("\n")
            if not cur.strip():
                break
            if re.match(r"^(#{1,4}\s|\||-{3,}$|[\s]*[-*]\s+)", cur):
                break
            para_lines.append(cur)
            i += 1
        yield {"type": "paragraph", "text": " ".join(para_lines)}


# -------------------------------------------------------------------
# DOCX writer
# -------------------------------------------------------------------


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _configure_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for s_name, size, bold, color in [
        ("Heading 1", 22, True, RGBColor(0x0A, 0x0A, 0x0A)),
        ("Heading 2", 18, True, RGBColor(0x0A, 0x0A, 0x0A)),
        ("Heading 3", 14, True, RGBColor(0x0A, 0x0A, 0x0A)),
        ("Heading 4", 12, True, RGBColor(0x52, 0x52, 0x52)),
    ]:
        h = doc.styles[s_name]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = color

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def render_block(doc: Document, block: dict) -> None:
    btype = block["type"]
    if btype == "heading":
        level = min(block["level"], 4)
        p = doc.add_paragraph(style=f"Heading {level}")
        _add_inline_runs(p, block["text"])
    elif btype == "paragraph":
        p = doc.add_paragraph()
        _add_inline_runs(p, block["text"])
    elif btype == "bullets":
        for item in block["items"]:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, item)
    elif btype == "hr":
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("• • •")
        run.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
    elif btype == "table":
        header = block["header"]
        rows = block["rows"]
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
        table.style = "Light Grid Accent 1"
        table.autofit = True

        for i, cell_text in enumerate(header):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.bold = True
            _set_cell_shading(cell, "F0F0F0")

        for r, row in enumerate(rows):
            for c, cell_text in enumerate(row):
                if c >= len(header):
                    continue
                cell = table.rows[r + 1].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_inline_runs(p, cell_text)

        doc.add_paragraph()


# -------------------------------------------------------------------
# Entry point
# -------------------------------------------------------------------


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines(keepends=True)
    doc = Document()
    _configure_document(doc)

    # Cover header
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title.add_run("neuthek")
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Operator legal-action checklist")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x52, 0x52, 0x52)

    doc.add_paragraph()

    for block in iter_blocks(lines):
        # Skip the source markdown's own H1 "neuthek — operator legal-
        # action checklist" since we already put it in the cover header.
        if block["type"] == "heading" and block["level"] == 1:
            continue
        render_block(doc, block)

    doc.save(str(docx_path))


if __name__ == "__main__":
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "LEGAL_CHECKLIST.md")
    dst = Path(sys.argv[2] if len(sys.argv) > 2 else "LEGAL_CHECKLIST.docx")
    md_to_docx(src, dst)
    print(f"OK: wrote {dst.absolute()} ({dst.stat().st_size:,} bytes)")
